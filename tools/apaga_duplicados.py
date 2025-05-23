from docx import Document

def remover_capitulos_duplicados(caminho_entrada, caminho_saida, caminho_log="log_duplicados.txt"):
    doc = Document(caminho_entrada)
    novo_doc = Document()

    titulos_vistos = set()
    capitulos_duplicados = []
    pular_capitulo = False

    for par in doc.paragraphs:
        if par.style.name == "Heading 2":
            titulo = par.text.strip()
            if titulo in titulos_vistos:
                capitulos_duplicados.append(titulo)
                print(f"⚠️ Capítulo duplicado detectado: {titulo}")
                pular_capitulo = True  # ativa o modo de pular até o próximo título
                continue
            else:
                titulos_vistos.add(titulo)
                novo_doc.add_paragraph(titulo, style="Heading 2")
                pular_capitulo = False  # volta ao modo normal
        else:
            if not pular_capitulo:
                novo_doc.add_paragraph(par.text, style=par.style)

    novo_doc.save(caminho_saida)
    print(f"✅ Novo arquivo salvo como: {caminho_saida}")

    if capitulos_duplicados:
        with open(caminho_log, "w", encoding="utf-8") as f:
            f.write("Capítulos duplicados removidos:\n\n")
            for titulo in capitulos_duplicados:
                f.write(f"{titulo}\n")
        print(f"📄 Log salvo como: {caminho_log}")
    else:
        print("✅ Nenhum capítulo duplicado encontrado.")



# Exemplo de uso
remover_capitulos_duplicados(
    r"F:\OneDrive\Documentos\Calibre\Roman Brocki\Divine Emperor of Death 1601-1900.docx",
    r"F:\OneDrive\Documentos\Calibre\Roman Brocki\Divine Emperor of Death 1601-1900_limpo.docx",
    "log_duplicados.txt"
)
