from docx import Document
import re
import os

from docx import Document
import re
import os

def desfazer_censura(texto):
    """
    Remove pontos de palavras que foram censuradas com letras separadas por '.'.
    Funciona para palavras inteiras e fragmentadas. Ex: s.e.x → sex, amus.ene.nt → amusement.
    Retorna: texto limpo e lista de palavras corrigidas
    """
    corrigidas = []

    def juntar_letras(match):
        original = match.group(0)
        corrigido = original.replace('.', '')
        if corrigido != original:
            corrigidas.append(original)
        return corrigido

    texto_corrigido = re.sub(r'\b(?:[a-zA-Z]\.?){3,}\b', juntar_letras, texto)
    return texto_corrigido, corrigidas

def limpar_censura_docx(caminho_docx):
    """
    Lê um .docx, remove censura por pontos, e salva como arquivo novo com sufixo '-uncens'.

    Parâmetro:
    - caminho_docx (str): caminho para o arquivo .docx original

    Retorna:
    - caminho_saida (str): caminho do novo arquivo salvo
    """
    doc = Document(caminho_docx)
    novo_doc = Document()
    log_corrigidas = []

    for par in doc.paragraphs:
        texto_corrigido, palavras = desfazer_censura(par.text)
        log_corrigidas.extend(palavras)
        novo_par = novo_doc.add_paragraph(texto_corrigido)
        novo_par.style = par.style

    nome_base, ext = os.path.splitext(caminho_docx)
    caminho_saida = f"{nome_base}_uncens{ext}"
    novo_doc.save(caminho_saida)

    # Gera log das palavras corrigidas
    if log_corrigidas:
        caminho_log = f"{nome_base}_uncens_log.txt"
        with open(caminho_log, "w", encoding="utf-8") as f:
            f.write("Palavras censuradas corrigidas:\n\n")
            for palavra in sorted(set(log_corrigidas)):
                f.write(f"{palavra} → {palavra.replace('.', '')}\n")
        print(f"📝 Log salvo em: {caminho_log}")

    return caminho_saida


limpar_censura_docx(r"F:\OneDrive\Documentos\Calibre\Roman Brocki\Divine Emperor of Death 2501-2800.docx")
