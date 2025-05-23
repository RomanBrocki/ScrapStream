#from selenium import webdriver
import os
#from selenium.webdriver.chrome.service import Service
#from selenium.webdriver.chrome.options import Options
import undetected_chromedriver as uc
import random
from selenium.webdriver.common.by import By  # Importando 'By' para usar os seletores
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from pattern import pattern, replacement  # Importando pattern e replacement para substituição no texto
from tools.desfaz_censura import desfazer_censura

import time
from config import config  # Importando o arquivo de configuração para os seletores

class NovelScraper:
    """
    A classe NovelScraper é responsável por realizar o scraping de capítulos de uma web novel,
    extrair o conteúdo da página e gerar um eBook no formato .docx.
    
    Funcionalidades principais:
    - Scraping de capítulos: Realiza a extração de conteúdo da web novel.
    - Substituição de padrões: Substitui padrões no texto do capítulo, conforme especificado no arquivo 'pattern.py'.
    - Geração de eBook: Compila o conteúdo extraído em um arquivo .docx.
    - Registro de ocorrências: Mantém um log das ocorrências de elementos específicos no conteúdo.
    """
    
    def __init__(self, config=config):
        """
        Inicializa a classe NovelScraper com um dicionário de configuração.

        Parâmetros:
        - config: Dicionário contendo as configurações dos seletores de elementos e outras informações, como o perfil do navegador.
        """
        self.config = config  # A configuração agora é passada diretamente para o método

        self.ebook_name = ""
        self.start_url = ""
        self.end_url = ""
        self.occurrence_list = []  # Lista para armazenar as ocorrências encontradas
        self.occurrences = 0  # Contagem total de ocorrências
        self.save_path = ""  # Caminho onde o eBook será salvo

    def scrape_chapters(self, start_url, end_url):
        """
        Realiza o scraping dos capítulos da web novel, extrai o conteúdo e cria um eBook no formato .docx.

        Parâmetros:
        - start_url: URL do primeiro capítulo da novel.
        - end_url: URL do último capítulo da novel.
        
        O método percorre os capítulos da novel, extrai o conteúdo e cria um arquivo .docx com o texto extraído.
        Ele também realiza a substituição de padrões definidos no arquivo 'pattern.py'.
        """
        self.start_url = start_url
        self.end_url = end_url

        #options = Options()
        #options.add_argument("user-data-dir=" + self.config['profile'])  # Usando o perfil configurado
        #browser = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        options = uc.ChromeOptions()
        options.add_argument(f"--user-data-dir={self.config['profile']}")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-blink-features=AutomationControlled")

        browser = uc.Chrome(options=options, headless=False)    

        document = Document()

        # Alterando o autor diretamente no documento, usando o valor de 'author' presente no config
        core_props = document.core_properties
        core_props.author = self.config['author']  # Define o autor

        # Lista para armazenar os títulos dos capítulos
        chapter_titles = []

        capitulo_count = 0             # Contador de capítulos
        falhas_recentes = 0            # Contador de falhas consecutivas
        modo_seguro = True             # Começa em modo seguro, com delays mais longos

        # começo do marcador de tempo
        inicio = time.time()
        
        while True:
            browser.get(self.start_url)

            # ⏳ Aguardar tempo aleatório para simular leitura real após carregar a página
            time.sleep(random.uniform(1.2, 2.2))

            # 🖱️ Scroll gradual para simular comportamento humano
            browser.execute_script("window.scrollTo(0, document.body.scrollHeight * 0.3);")
            time.sleep(random.uniform(0.4, 0.7))
            browser.execute_script("window.scrollTo(0, document.body.scrollHeight * 0.7);")
            time.sleep(random.uniform(0.4, 0.7))
            browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")

            try:
                # Espera até que o título do capítulo e o conteúdo estejam carregados na página
                WebDriverWait(browser, 20).until(EC.presence_of_element_located(self.config['chapter_title_selector']))
                WebDriverWait(browser, 20).until(EC.presence_of_element_located(self.config['chapter_content_selector']))

                # Extrai o título e o conteúdo do capítulo
                chapter_title = browser.find_element(*self.config['chapter_title_selector']).text
                chapter_content = browser.find_element(*self.config['chapter_content_selector']).text
                
                # Remove censura por pontos
                chapter_content = desfazer_censura(chapter_content)


                # ✅ Sucesso: reseta contador de falhas
                falhas_recentes = 0

            except Exception as e:
                print(f"Error loading page or elements not found. Retrying... Error: {e}")
                falhas_recentes += 1
                if falhas_recentes >= 2:
                    modo_seguro = True  # Ativa modo seguro após 2 falhas seguidas
                time.sleep(5)
                continue  # Tenta novamente caso o erro ocorra

            # Verifica se o capítulo já foi copiado para evitar duplicatas
            capitulo_duplicado = chapter_title in chapter_titles

            if capitulo_duplicado:
                print(f"⚠️ Capítulo repetido detectado: '{chapter_title}'. Conteúdo será ignorado, mas avançando normalmente.")
            else:
                # Adiciona o título do capítulo à lista de capítulos
                chapter_titles.append(chapter_title)

                # Substitui os padrões definidos em `pattern.py`
                for element in pattern:
                    if element in chapter_content:
                        self.occurrence_list.append(f'Found: {element}')  # Registra a ocorrência
                        self.occurrences += 1  # Incrementa o contador de ocorrências
                    chapter_content = chapter_content.replace(element, replacement)  # Substitui o padrão encontrado

                # Adiciona o capítulo ao documento
                document.add_paragraph(f'{chapter_title}', style='Heading 2')  # Título do capítulo
                document.add_paragraph(chapter_content)  # Conteúdo do capítulo
                document.add_page_break()


            if self.start_url == self.end_url:
                break  # Se o scraping chegou ao último capítulo, sai do loop

            try:
                # Espera até que o botão de próximo capítulo esteja disponível e clica nele
                WebDriverWait(browser, 20).until(EC.presence_of_element_located(self.config['next_chapter_selector']))
                browser.find_element(*self.config['next_chapter_selector']).click()
            except Exception as e:
                print(f"Next chapter button not found or error clicking next: {e}")
                break  # Se não encontrar o próximo capítulo ou não conseguir clicar, interrompe o processo

            # Atualiza a URL para o próximo capítulo
            self.start_url = browser.current_url

            # 🔁 Pausa adaptativa entre capítulos
            capitulo_count += 1

            if modo_seguro:
                if capitulo_count % 5 == 0:
                    time.sleep(random.uniform(5.0, 8.0))  # Pausa maior no modo seguro
                else:
                    time.sleep(random.uniform(1.0, 2.0))  # Pausa padrão no modo seguro

                # 🧠 Se passou 20 capítulos sem erro, volta ao modo rápido
                if falhas_recentes == 0 and capitulo_count % 20 == 0:
                    print("✅ Estável há 20 capítulos. Retornando ao modo rápido.")
                    modo_seguro = False
            else:
                if capitulo_count % 5 == 0:
                    time.sleep(random.uniform(2.0, 3.0))  # Pausa leve no modo rápido
                else:
                    time.sleep(random.uniform(0.4, 1.0))  # Pausa curta no modo rápido

        # --- REMOVIDO: índice manual (texto simples) que não é funcional no Kindle ---
        # Este trecho adicionava um índice de capítulos ao final do documento como parágrafos normais.
        # Como o Calibre usa os estilos (Heading 2, etc.) para gerar um índice clicável, 
        # este bloco não é mais necessário e pode ser desativado.
        
        # Inserir o índice de capítulos no início do documento
        # document.add_paragraph('Table of Contents', style='Heading 1')  # Título do índice
        # for title in chapter_titles:
        #     document.add_paragraph(f'{title}', style='Normal')  # Adiciona cada capítulo ao índice

        # Salva o documento somente após todos os capítulos e o índice terem sido adicionados
        
        # Calculo da marcação de tempo
        fim = time.time()
        duracao_segundos = fim - inicio
        duracao_minutos = duracao_segundos / 60
        media_por_capitulo = duracao_segundos / capitulo_count if capitulo_count else 0

        print(f"📘 Scraping finalizado.")
        print(f"⏱️ Duração total: {duracao_minutos:.2f} minutos")
        print(f"📄 Capítulos raspados: {capitulo_count}")
        print(f"🕒 Média por capítulo: {media_por_capitulo:.2f} segundos")
        


        document.save(os.path.join(self.save_path, "Novel.docx"))

        browser.quit()

        return {
            "duracao_min": round(duracao_minutos, 2),
            "capitulos": capitulo_count,
            "media_seg": round(media_por_capitulo, 2)
        }



    def get_log(self):
        """
        Retorna um log com todas as ocorrências encontradas durante o scraping.
        
        O log inclui todos os padrões encontrados no conteúdo e o número total de ocorrências.
        """
        log_list = ['List of found occurrences:']
        for occurrence in self.occurrence_list:
            log_list.append(occurrence)  # Adiciona cada ocorrência ao log
        log_list.append(f'Total occurrences found: {self.occurrences}')  # Adiciona o total de ocorrências
        return log_list

    def save_ebook(self, save_path, ebook_name):
        """
        Renomeia e salva o arquivo .docx com o nome do eBook fornecido.

        Parâmetros:
        - save_path: Caminho onde o eBook será salvo.
        - ebook_name: Nome do eBook a ser atribuído ao arquivo .docx.
        """
        self.ebook_name = ebook_name
        self.save_path = save_path
        os.rename(os.path.join(self.save_path, "Novel.docx"), os.path.join(self.save_path, f"{self.ebook_name}.docx"))

    def run_scraper(self, ebook_name, save_path, start_url, end_url):
        """
        Executa o processo completo de scraping e salva o eBook gerado.

        Parâmetros:
        - ebook_name: Nome do eBook a ser gerado.
        - save_path: Caminho onde o eBook será salvo.
        - start_url: URL do primeiro capítulo.
        - end_url: URL do último capítulo.
        """
        self.save_path = save_path
        stats = self.scrape_chapters(start_url, end_url)  # Realiza o scraping dos capítulos e pega o retorno da mensuração de tempo
        self.save_ebook(save_path, ebook_name)  # Salva o eBook gerado
        return stats













